from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# File paths
txt_path = "docs/chat.txt"
docx_path = "docs/chat.docx"
user_name = "Arturo-Quiroga-MSFT"

def is_question(line):
    return line.strip().startswith(f"{user_name}:") and "?" in line

def add_bold_run(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return run

def main():
    doc = Document()
    with open(txt_path, "r", encoding="utf-8") as f:
        for line in f:
            if is_question(line):
                p = doc.add_paragraph()
                add_bold_run(p, line.strip())
            else:
                doc.add_paragraph(line.rstrip())
    doc.save(docx_path)
    print(f"Saved as {docx_path}")

if __name__ == "__main__":
    main()